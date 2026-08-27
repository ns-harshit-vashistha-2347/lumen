"""Document parsers.

Beyond just extracting text, we lift out:
- Tables (rendered as Markdown pipe-tables so the LLM keeps row/column structure)
- Images (extracted, optionally OCR'd, referenced by path in metadata)

OCR uses pytesseract if the `tesseract` binary is installed on the host.
If it isn't, images fall back to a placeholder caption so downstream
retrieval can still find them by page/position.
"""
from __future__ import annotations

import io
import os
from collections import Counter
from pathlib import Path
from typing import Iterable

import pdfplumber
from docx import Document as DocxDocument

from src.core.config import settings
from src.core.logging import get_logger
from src.interfaces.base_parser import BaseParser, ParsedUnit

logger = get_logger(__name__)


# --- shared helpers ---------------------------------------------------------


def _images_dir_for(file_path: str) -> Path:
    base = Path(getattr(settings, "UPLOAD_DIR", ".")) / "images"
    stem = Path(file_path).stem
    out = base / stem
    out.mkdir(parents=True, exist_ok=True)
    return out


def _ocr_available() -> bool:
    if not getattr(settings, "OCR_ENABLED", True):
        return False
    try:
        import pytesseract  # noqa: F401
        from PIL import Image  # noqa: F401
        return True
    except ImportError:
        return False


def _ocr_image_bytes(data: bytes) -> str:
    """Best-effort OCR. Returns '' on any failure."""
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(io.BytesIO(data))
        text = pytesseract.image_to_string(img) or ""
        return text.strip()
    except Exception as exc:
        logger.debug(f"[parse.ocr] failed: {exc}")
        return ""


def _table_to_markdown(rows: Iterable[Iterable[object]]) -> str:
    """Render a 2D list as a GitHub-style markdown table. Empty cells become
    a single dash so downstream chunkers don't collapse the column layout."""
    normalized = [[str(c or "").strip() or "-" for c in row] for row in rows]
    normalized = [r for r in normalized if any(c and c != "-" for c in r)]
    if not normalized:
        return ""
    header = normalized[0]
    body = normalized[1:] if len(normalized) > 1 else []
    width = len(header)
    body = [(r + [""] * (width - len(r)))[:width] for r in body]

    md_lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * width) + " |",
    ]
    md_lines.extend("| " + " | ".join(r) + " |" for r in body)
    return "\n".join(md_lines)


# --- PDF --------------------------------------------------------------------


class PdfParser(BaseParser):
    def supports(self, file_path: str) -> bool:
        return file_path.lower().endswith(".pdf")

    def parse(self, file_path: str) -> list[ParsedUnit]:
        with pdfplumber.open(file_path) as pdf:
            doc_meta = pdf.metadata or {}
            pages_text: list[str] = []
            per_page_tables: list[list[str]] = []
            for page in pdf.pages:
                pages_text.append(page.extract_text() or "")
                page_tables: list[str] = []
                try:
                    for tbl in page.extract_tables() or []:
                        md = _table_to_markdown(tbl)
                        if md:
                            page_tables.append(md)
                except Exception as exc:
                    logger.debug(f"[pdf] table extract failed p={page.page_number}: {exc}")
                per_page_tables.append(page_tables)

        pages_text = self._strip_repeated_headers_footers(pages_text)

        units: list[ParsedUnit] = []
        source_name = os.path.basename(file_path)
        doc_title = doc_meta.get("Title") or source_name
        doc_author = doc_meta.get("Author")

        for page_num, text in enumerate(pages_text, start=1):
            body_parts: list[str] = []
            if text.strip():
                body_parts.append(text.strip())
            for tbl_md in per_page_tables[page_num - 1]:
                body_parts.append(f"\n\n[TABLE]\n{tbl_md}\n[/TABLE]\n")

            merged = "\n\n".join(body_parts).strip()
            if merged:
                units.append(ParsedUnit(
                    content=merged,
                    metadata={
                        "page_number": page_num,
                        "source": source_name,
                        "doc_title": doc_title,
                        "doc_author": doc_author,
                        "content_kind": "text",
                    },
                ))

        # Images: emit each as its own unit with OCR'd or placeholder caption.
        try:
            units.extend(self._extract_pdf_images(file_path, source_name))
        except Exception as exc:
            logger.warning(f"[pdf] image extraction failed for {source_name}: {exc}")

        return units

    def _extract_pdf_images(self, file_path: str, source_name: str) -> list[ParsedUnit]:
        try:
            import pypdf
        except ImportError:
            return []

        out: list[ParsedUnit] = []
        ocr_on = _ocr_available()
        images_dir = _images_dir_for(file_path)

        reader = pypdf.PdfReader(file_path)
        for page_num, page in enumerate(reader.pages, start=1):
            try:
                images = getattr(page, "images", []) or []
            except Exception:
                images = []
            for img_ix, img in enumerate(images, start=1):
                try:
                    data = img.data
                    ext = (img.name.split(".")[-1] if "." in img.name else "png").lower()
                except Exception:
                    continue
                out_path = images_dir / f"page{page_num}_img{img_ix}.{ext}"
                try:
                    with open(out_path, "wb") as fh:
                        fh.write(data)
                except OSError as exc:
                    logger.debug(f"[pdf.image] failed to write {out_path}: {exc}")
                    continue

                caption = _ocr_image_bytes(data) if ocr_on else ""
                if not caption:
                    caption = f"[Image on page {page_num}, position {img_ix}. OCR unavailable or empty.]"
                out.append(ParsedUnit(
                    content=f"[IMAGE p{page_num} #{img_ix}]\n{caption}",
                    metadata={
                        "page_number": page_num,
                        "source": source_name,
                        "content_kind": "image",
                        "image_path": str(out_path),
                        "image_ocr": bool(caption and not caption.startswith("[Image")),
                    },
                ))
        return out

    def _strip_repeated_headers_footers(self, pages: list[str], min_pages: int = 3) -> list[str]:
        """Lines repeated across most pages are running headers/footers — drop them."""
        if len(pages) < min_pages:
            return pages
        first_lines = Counter(p.split("\n", 1)[0].strip() for p in pages if p.strip())
        last_lines = Counter(p.rstrip().rsplit("\n", 1)[-1].strip() for p in pages if p.strip())
        threshold = len(pages) * 0.6
        noisy = {ln for ln, c in first_lines.items() if c >= threshold and ln} | \
                {ln for ln, c in last_lines.items() if c >= threshold and ln}

        cleaned = []
        for p in pages:
            lines = [ln for ln in p.split("\n") if ln.strip() not in noisy]
            cleaned.append("\n".join(lines))
        return cleaned


# --- DOCX -------------------------------------------------------------------


class DocxParser(BaseParser):
    def supports(self, file_path: str) -> bool:
        return file_path.lower().endswith(".docx")

    def parse(self, file_path: str) -> list[ParsedUnit]:
        doc = DocxDocument(file_path)
        source_name = os.path.basename(file_path)
        units: list[ParsedUnit] = []

        # Body paragraphs
        body = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if body.strip():
            units.append(ParsedUnit(
                content=body,
                metadata={"source": source_name, "content_kind": "text"},
            ))

        # Tables as markdown
        for tbl_ix, table in enumerate(doc.tables, start=1):
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            md = _table_to_markdown(rows)
            if md:
                units.append(ParsedUnit(
                    content=f"[TABLE #{tbl_ix}]\n{md}\n[/TABLE]",
                    metadata={
                        "source": source_name,
                        "content_kind": "table",
                        "table_index": tbl_ix,
                    },
                ))

        # Images embedded in the docx package
        try:
            units.extend(self._extract_docx_images(doc, file_path, source_name))
        except Exception as exc:
            logger.warning(f"[docx] image extraction failed for {source_name}: {exc}")

        return units

    def _extract_docx_images(self, doc, file_path: str, source_name: str) -> list[ParsedUnit]:
        out: list[ParsedUnit] = []
        images_dir = _images_dir_for(file_path)
        ocr_on = _ocr_available()

        # docx.package.image_parts is the reliable API across versions.
        try:
            image_parts = list(doc.part.package.image_parts)
        except Exception:
            image_parts = []

        for ix, part in enumerate(image_parts, start=1):
            try:
                data = part.blob
                ext = (part.partname.split(".")[-1] or "png").lower()
            except Exception:
                continue
            out_path = images_dir / f"img{ix}.{ext}"
            try:
                with open(out_path, "wb") as fh:
                    fh.write(data)
            except OSError as exc:
                logger.debug(f"[docx.image] failed to write {out_path}: {exc}")
                continue
            caption = _ocr_image_bytes(data) if ocr_on else ""
            if not caption:
                caption = f"[Image #{ix} in {source_name}. OCR unavailable or empty.]"
            out.append(ParsedUnit(
                content=f"[IMAGE #{ix}]\n{caption}",
                metadata={
                    "source": source_name,
                    "content_kind": "image",
                    "image_path": str(out_path),
                    "image_ocr": bool(caption and not caption.startswith("[Image")),
                },
            ))
        return out


class MarkdownTextParser(BaseParser):
    def supports(self, file_path: str) -> bool:
        return file_path.lower().endswith((".md", ".txt"))

    def parse(self, file_path: str) -> list[ParsedUnit]:
        with open(file_path, encoding="utf-8") as f:
            text = f.read()
        return [ParsedUnit(
            content=text,
            metadata={"source": os.path.basename(file_path), "content_kind": "text"},
        )]


DOCUMENT_PARSERS: list[BaseParser] = [
    PdfParser(),
    DocxParser(),
    MarkdownTextParser(),
]


def get_parser(file_path: str) -> BaseParser:
    for parser in DOCUMENT_PARSERS:
        if parser.supports(file_path):
            return parser
    raise ValueError(f"No parser registered for file: {file_path}")


def parse_node(state: dict) -> dict:
    file_path = state.get("file_path")
    if not file_path:
        raise ValueError("file_path is required in the state.")

    logger.info(f"Parsing file: {file_path}")
    parser = get_parser(file_path)
    units = parser.parse(file_path)
    if not units:
        raise ValueError(f"No content extracted from file: {file_path}")

    counts = Counter((u.metadata.get("content_kind", "text")) for u in units)
    logger.info(f"[parse_node] {file_path}: {dict(counts)}")
    return {**state, "parsed_units": units}
